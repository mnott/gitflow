"""
This module provides functionality for generating documentation from GitHub issues.
"""

from typing import List, Dict, Optional, Set
import re
from dataclasses import dataclass
import json
import subprocess
from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn, TaskProgressColumn
from rich.console import Console
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
from bs4 import BeautifulSoup
from docx.oxml.shared import OxmlElement, qn
from .docstyles import apply_styles_to_document
import os
import docx.opc
import docx.opc.constants
import requests
from io import BytesIO
from urllib.parse import urlparse
import urllib3
import certifi
import tempfile
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

console = Console()

@dataclass
class IssueContent:
    """Represents the content of a GitHub issue"""
    number: int
    title: str
    description: str
    children: List[int]  # List of child issue numbers
    comments: List[Dict] = None  # Optional list of comments

class IssueDocGenerator:
    """
    A class for generating documentation from GitHub issues.
    """

    def __init__(self):
        """Initialize the IssueDocGenerator."""
        self.issues_cache: Dict[int, IssueContent] = {}
        self.processed_issues: Set[int] = set()
        self.bookmarks: Dict[int, str] = {}  # Track bookmarks for each issue

        # Get repo URL for links
        try:
            result = subprocess.run(
                ["gh", "repo", "view", "--json", "url"],
                capture_output=True, text=True, check=True
            )
            self.repo_url = json.loads(result.stdout)['url']
        except subprocess.CalledProcessError:
            console.print("[yellow]Warning: Could not get repository URL. Links will be disabled.[/yellow]")
            self.repo_url = None

        # Get path to styles.css relative to this module
        module_dir = os.path.dirname(os.path.abspath(__file__))
        self.css_path = os.path.join(module_dir, 'styles.css')
        console.print(f"[blue]Looking for styles.css at: {self.css_path}[/blue]")

        if not os.path.exists(self.css_path):
            console.print(f"[red]Warning: Could not find styles.css at {self.css_path}[/red]")
        else:
            console.print(f"[green]Found styles.css at {self.css_path}[/green]")

    def _extract_issue_numbers(self, text: str) -> List[int]:
        """Extract issue numbers from text using regex."""
        matches = re.finditer(r'#(\d+)', text)
        return [int(match.group(1)) for match in matches]

    def _get_issue(self, number: int, progress, task_id, include_comments: bool = False) -> Optional[IssueContent]:
        """Fetch issue content using gh CLI."""
        if number in self.issues_cache:
            return self.issues_cache[number]

        try:
            progress.update(task_id, description=f"[blue]Fetching issue #{number}...")

            # Fetch basic issue data
            result = subprocess.run(
                ["gh", "issue", "view", str(number), "--json", "number,title,body"],
                capture_output=True, text=True, check=True
            )
            data = json.loads(result.stdout)
            children = self._extract_issue_numbers(data['body'])

            # Fetch comments if requested
            comments = None
            if include_comments:
                result = subprocess.run(
                    ["gh", "issue", "view", str(number), "--json", "comments"],
                    capture_output=True, text=True, check=True
                )
                comments_data = json.loads(result.stdout)
                comments = comments_data.get('comments', [])

            issue = IssueContent(
                number=data['number'],
                title=data['title'],
                description=data['body'],
                children=children,
                comments=comments
            )

            self.issues_cache[number] = issue
            return issue

        except subprocess.CalledProcessError:
            console.print(f"[red]Failed to fetch issue #{number}[/red]")
            return None

    def _apply_heading_style(self, paragraph, level):
        """Apply heading styles from CSS."""
        # Map level to style name
        if level == 1:
            paragraph.style = 'Heading 1'
        elif level == 2:
            paragraph.style = 'Heading 2'
        elif level == 3:
            paragraph.style = 'Heading 3'
        elif level == 4:
            paragraph.style = 'Heading 4'
        else:
            paragraph.style = f'Heading {level}'

    def _add_bookmark(self, paragraph, issue_num: int):
        """Add a bookmark to a paragraph for internal linking."""
        bookmark = OxmlElement('w:bookmarkStart')
        bookmark_id = f'issue_{issue_num}'
        bookmark.set(qn('w:id'), '0')
        bookmark.set(qn('w:name'), bookmark_id)
        paragraph._p.append(bookmark)

        # Add bookmark end
        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), '0')
        paragraph._p.append(bookmark_end)

        self.bookmarks[issue_num] = bookmark_id

    def _create_internal_hyperlink(self, paragraph, text: str, bookmark_id: str, is_bullet: bool = True):
        """Create a hyperlink to a bookmark within the document."""
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), bookmark_id)

        # Create run element with specific hyperlink style
        run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Style the link
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'Hyperlink')  # Use Word's built-in hyperlink style
        rPr.append(rStyle)

        # Set font size for bullet points
        if is_bullet and self.doc_styles.bullet_font_size:
            sz = OxmlElement('w:sz')
            size = self.doc_styles.bullet_font_size
            word_units = size * 2
            sz.set(qn('w:val'), str(word_units))
            rPr.append(sz)

        # Add text
        t = OxmlElement('w:t')
        t.text = text
        run.append(rPr)
        run.append(t)

        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    def _create_external_hyperlink(self, paragraph, text: str, url: str, is_bullet: bool = True):
        """Create a hyperlink to an external URL."""
        # Create the hyperlink run
        run = paragraph.add_run()
        run.font.underline = True
        run.font.color.rgb = RGBColor(5, 99, 193)  # Office blue
        if is_bullet and self.doc_styles.bullet_font_size:
            run.font.size = Pt(self.doc_styles.bullet_font_size)

        # Get the XML element
        r = run._r

        # Create hyperlink element
        hyperlink = OxmlElement('w:hyperlink')

        # Create the relationship
        r_id = paragraph.part.relate_to(
            url,
            docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
            is_external=True
        )
        hyperlink.set(qn('r:id'), r_id)

        # Add text
        t = OxmlElement('w:t')
        t.text = text
        r.append(t)

        # Replace run with hyperlink containing the run
        paragraph._p.append(hyperlink)
        hyperlink.append(r)

    def _apply_body_style(self, paragraph, text, is_bullet: bool = False):
        """Apply body text styles and handle links."""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # First handle internal links
        parts = re.split(r'(\{INTERNAL_LINK:\d+:.*?\})', text)

        for part in parts:
            if part.startswith('{INTERNAL_LINK:'):
                # Extract issue number and title
                match = re.match(r'\{INTERNAL_LINK:(\d+):(.*?)\}', part)
                if match:
                    issue_num = int(match.group(1))
                    title = match.group(2)
                    if issue_num in self.bookmarks:
                        self._create_internal_hyperlink(
                            paragraph,
                            f"#{issue_num} - {title}",
                            self.bookmarks[issue_num],
                            is_bullet=is_bullet
                        )
                    else:
                        # If we don't have the bookmark, just show as plain text
                        run = paragraph.add_run(f"#{issue_num} - {title}")
                        if is_bullet:
                            run.font.size = Pt(self.doc_styles.bullet_font_size)
                        run.font.color.rgb = RGBColor(128, 128, 128)
            else:
                # Convert to HTML to handle links
                html = markdown.markdown(part)
                soup = BeautifulSoup(html, 'html.parser')

                # Process text with links
                current_pos = 0
                text = soup.get_text()

                # Find all links and their positions
                links = []
                for link in soup.find_all('a'):
                    url = link.get('href', '')
                    link_text = link.get_text()
                    # Find position in the text
                    start = text.find(link_text, current_pos)
                    if start != -1:
                        links.append({
                            'url': url,
                            'text': link_text,
                            'position': start
                        })
                        current_pos = start + len(link_text)

                # Add text with links
                last_pos = 0
                for link in sorted(links, key=lambda x: x['position']):
                    # Add text before link
                    if link['position'] > last_pos:
                        pre_link_text = text[last_pos:link['position']]
                        if pre_link_text:
                            run = paragraph.add_run(pre_link_text)
                            if is_bullet:
                                run.font.size = Pt(self.doc_styles.bullet_font_size)
                    # Add link
                    self._create_external_hyperlink(paragraph, link['text'], link['url'], is_bullet)
                    last_pos = link['position'] + len(link['text'])

                # Add remaining text
                if last_pos < len(text):
                    remaining_text = text[last_pos:]
                    if remaining_text:
                        run = paragraph.add_run(remaining_text)
                        if is_bullet:
                            run.font.size = Pt(self.doc_styles.bullet_font_size)

    def _apply_code_style(self, paragraph, text):
        """Apply code block styling."""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Get code block style properties
        code_props = self.doc_styles.styles.get('codeblock', {})

        # Create and style the run
        run = paragraph.add_run(text)

        # Apply font family
        run.font.name = code_props.get('font-family', 'Consolas').strip("'")

        # Apply font size
        if 'font-size' in code_props:
            size = int(re.match(r'(\d+)px', code_props['font-size']).group(1))
            run.font.size = Pt(size)

        # Apply background color
        if 'background-color' in code_props:
            color = code_props['background-color'].strip('#')
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), color)
            paragraph._element.get_or_add_pPr().append(shading_elm)

        # Apply margins and padding
        pPr = paragraph._element.get_or_add_pPr()

        if 'margin-left' in code_props:
            margin = int(re.match(r'(\d+)px', code_props['margin-left']).group(1))
            ind = pPr.get_or_add_ind()
            ind.set(qn('w:left'), str(margin * 20))

        if 'padding' in code_props:
            padding = int(re.match(r'(\d+)px', code_props['padding']).group(1))
            spacing = pPr.get_or_add_spacing()
            spacing.set(qn('w:before'), str(padding * 20))
            spacing.set(qn('w:after'), str(padding * 20))

        # Apply line spacing
        if 'line-height' in code_props:
            height = int(re.match(r'(\d+)px', code_props['line-height']).group(1))
            spacing = pPr.get_or_add_spacing()
            spacing.set(qn('w:line'), str(height * 20))
            spacing.set(qn('w:lineRule'), 'exact')

    def _process_issue_content(self, content: str, depth: int = 0) -> str:
        """Process issue content, replacing issue numbers with titles."""
        def replace_issue_ref(match):
            issue_num = int(match.group(1))
            if issue_num in self.issues_cache:
                title = self.issues_cache[issue_num].title
                return f"{{INTERNAL_LINK:{issue_num}:{title}}}"
            return f"-{issue_num}"

        # Process issue references
        content = re.sub(r'#(\d+)', replace_issue_ref, content)
        return content

    def generate_doc(self, root_issue: int, max_depth: int = 1, max_issues: Optional[int] = None,
                    include_comments: bool = False) -> Document:
        """Generate a Word document from the issue tree."""
        doc = Document()
        self.processed_count = 0

        # Apply styles from CSS and store the style object
        self.doc_styles = apply_styles_to_document(doc, self.css_path)

        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            BarColumn(),
            TaskProgressColumn(),
            console=console
        ) as progress:
            # First, fetch all issues and create bookmarks
            fetch_task = progress.add_task("[blue]Fetching issues...", total=None)
            self._fetch_all_issues(root_issue, max_depth, max_issues, progress, fetch_task, include_comments)

            # Then process the content
            process_task = progress.add_task("[blue]Processing content...", total=None)
            self._process_issue_tree(root_issue, doc, 0, max_depth, progress, process_task, max_issues, include_comments)
            progress.update(process_task, completed=True)

        return doc

    def _fetch_all_issues(self, root_issue: int, max_depth: int, max_issues: Optional[int],
                         progress, task_id, include_comments: bool = False):
        """Fetch all issues first to ensure we have complete data."""
        issues_to_fetch = [(root_issue, 0)]  # (issue_number, depth)
        fetched_issues = set()

        while issues_to_fetch and (max_issues is None or len(fetched_issues) < max_issues):
            current_issue, current_depth = issues_to_fetch.pop(0)  # Use FIFO order
            if current_issue not in fetched_issues:
                issue = self._get_issue(current_issue, progress, task_id, include_comments)
                if issue:
                    fetched_issues.add(current_issue)
                    # Create a temporary paragraph to add bookmark
                    self.bookmarks[current_issue] = f'issue_{current_issue}'

                    # Add child issues to fetch queue if within depth
                    if current_depth < max_depth:
                        for child in issue.children:
                            if child not in fetched_issues:
                                issues_to_fetch.append((child, current_depth + 1))

        progress.update(task_id, description=f"[blue]Fetched {len(fetched_issues)} issues")

    def _add_image(self, doc, image_url: str):
        """Download and add image to document."""
        try:
            # Handle GitHub images that require authentication
            if 'github' in image_url:
                # Create a temporary file that will be automatically cleaned up
                with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                    # Use wget exactly as it works in terminal
                    result = subprocess.run(
                        ["wget", "-O", temp_file.name, "--no-check-certificate", image_url],
                        capture_output=True, check=False
                    )

                    if result.returncode != 0:
                        raise Exception(f"wget command failed: {result.stderr.decode()}")

                    # Read the downloaded file
                    with open(temp_file.name, 'rb') as f:
                        image_data = f.read()

                    # Clean up temp file
                    os.unlink(temp_file.name)

                    image_stream = BytesIO(image_data)
            else:
                response = requests.get(image_url)
                response.raise_for_status()
                image_stream = BytesIO(response.content)

            # Add image in a table cell to create a frame
            table = doc.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)

            # Add the image to the cell
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(image_stream, width=Inches(6.0))

            # Style the table/frame
            table.style = 'Table Grid'
            table.allow_autofit = False

            # Add cell padding (margins)
            for row in table.rows:
                for cell in row.cells:
                    cell.width = Inches(6.5)  # Slightly wider than image
                    paragraph = cell.paragraphs[0]
                    paragraph.paragraph_format.space_before = Pt(12)
                    paragraph.paragraph_format.space_after = Pt(12)
                    paragraph.paragraph_format.left_indent = Pt(12)
                    paragraph.paragraph_format.right_indent = Pt(12)

            # Add spacing after the table
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(12)

        except Exception as e:
            console.print(f"[red]Failed to add image from {image_url}: {str(e)}[/red]")
            # Add a placeholder text for failed images
            para = doc.add_paragraph(f"[Image could not be loaded: {image_url}]")
            para.italic = True
            para.paragraph_format.space_after = Pt(12)

    def _process_issue_tree(self, issue_num: int, doc: Document, current_depth: int, max_depth: int,
                           progress, task_id, max_issues: Optional[int] = None, include_comments: bool = False):
        """Recursively process issues and add content to document."""
        if current_depth > max_depth or issue_num in self.processed_issues:
            return

        if max_issues and self.processed_count >= max_issues:
            return

        issue = self.issues_cache.get(issue_num)
        if not issue:
            return

        self.processed_issues.add(issue_num)
        self.processed_count += 1
        progress.update(task_id, description=f"[blue]Processing issue #{issue_num} ({self.processed_count} issues processed)...")

        # Add a page break before each new issue (except the first one)
        if self.processed_count > 1:
            doc.add_page_break()

        # Add issue title with proper styling and bookmark
        is_top_level = current_depth == 0
        if is_top_level and self.processed_count > 1:
            heading_level = 1
        else:
            heading_level = current_depth + 1

        heading = doc.add_heading(f"#{issue_num} - {issue.title}", level=heading_level)
        self._apply_heading_style(heading, heading_level)
        self._add_bookmark(heading, issue_num)

        # First process the description to replace issue references
        processed_description = self._process_issue_content(issue.description, current_depth)

        # Then convert to HTML with headers extension
        html = markdown.markdown(processed_description,
                               extensions=['fenced_code', 'codehilite', 'markdown.extensions.toc'])
        soup = BeautifulSoup(html, 'html.parser')

        def process_element_with_links(element, paragraph, is_bullet=False):
            """Helper function to process text with links consistently."""
            # Get the raw HTML content to preserve links
            html_content = str(element)
            item_soup = BeautifulSoup(html_content, 'html.parser')

            # Find all links and their positions
            links = []
            for link in item_soup.find_all('a'):
                # Get the text before this link
                text_before = ''.join(t.string or '' for t in link.previous_siblings)
                links.append({
                    'url': link.get('href', ''),
                    'text': link.get_text(),
                    'position': len(text_before)
                })
                # Remove the link but keep its text
                link.unwrap()

            # Process the full text
            text = self._process_issue_content(item_soup.get_text().strip(), current_depth)
            if text.strip():
                # Split text at link positions and add pieces with links
                last_pos = 0
                for link in sorted(links, key=lambda x: x['position']):
                    # Add text before link
                    if link['position'] > last_pos:
                        self._apply_body_style(paragraph, text[last_pos:link['position']], is_bullet=is_bullet)
                    # Add link
                    self._create_external_hyperlink(paragraph, link['text'], link['url'], is_bullet=is_bullet)
                    last_pos = link['position'] + len(link['text'])
                # Add remaining text
                if last_pos < len(text):
                    self._apply_body_style(paragraph, text[last_pos:], is_bullet=is_bullet)

        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'pre', 'img']):
            if element.name == 'img':
                # Handle image
                image_url = element.get('src', '')
                if image_url:
                    self._add_image(doc, image_url)
            elif element.name.startswith('h'):
                level = int(element.name[1])  # Get number from h1, h2, etc.
                # For subsequent documents at top level, don't add current_depth
                if is_top_level and self.processed_count > 1:
                    effective_level = min(level, 4)
                else:
                    effective_level = min(current_depth + level, 4)
                heading = doc.add_heading(element.get_text(), level=effective_level)
                self._apply_heading_style(heading, effective_level)
            elif element.name == 'pre':
                # Handle code blocks
                paragraph = doc.add_paragraph()
                self._apply_code_style(paragraph, element.get_text())
            elif element.name == 'ul':
                for li in element.find_all('li'):
                    paragraph = doc.add_paragraph(style='List Bullet')
                    process_element_with_links(li, paragraph, is_bullet=True)
            elif element.name == 'p':
                # Handle paragraphs with potential links
                paragraph = doc.add_paragraph()
                process_element_with_links(element, paragraph, is_bullet=False)

        # Add comments if requested
        if include_comments and issue.comments:
            comments_heading = doc.add_heading('Comments', level=min(current_depth + 2, 4))
            self._apply_heading_style(comments_heading, min(current_depth + 2, 4))

            for comment in issue.comments:
                # Convert comment body to HTML
                comment_html = markdown.markdown(comment['body'],
                                              extensions=['fenced_code', 'codehilite'])
                comment_soup = BeautifulSoup(comment_html, 'html.parser')

                # Process each element in the comment
                for comment_element in comment_soup.find_all(['p', 'pre', 'ul']):
                    if comment_element.name == 'pre':
                        paragraph = doc.add_paragraph()
                        self._apply_code_style(paragraph, comment_element.get_text())
                    elif comment_element.name == 'ul':
                        for li in comment_element.find_all('li'):
                            paragraph = doc.add_paragraph(style='List Bullet')
                            process_element_with_links(li, paragraph, is_bullet=True)
                    else:
                        paragraph = doc.add_paragraph()
                        process_element_with_links(comment_element, paragraph, is_bullet=False)

        # Process child issues if within depth limit
        if current_depth < max_depth:
            for child_num in issue.children:
                if child_num not in self.processed_issues:
                    self._process_issue_tree(child_num, doc, current_depth + 1, max_depth,
                                          progress, task_id, max_issues, include_comments)

    def _get_relationship_id(self, part, url: str) -> str:
        """Get or create relationship ID for external URL."""
        # Check if relationship already exists
        for rel in part.rels:
            if part.rels[rel].target_ref == url:
                return rel

        # Create new relationship
        rel_id = f'rId{len(part.rels) + 1}'
        part.rels.add_relationship(
            rel_id,
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
            url,
            'External'
        )
        return rel_id